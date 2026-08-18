#!/usr/bin/env python3
"""Convert the checked enzyme workbook into the lecture 12 site payload."""

from __future__ import annotations

import json
import random
import re
from pathlib import Path

from docx import Document


SOURCE = Path("/Users/ray/Downloads/生化_酶_学成选择题_优化打散版.docx")
LECTURE_NUMBER = 12
TITLE = "生化 酶"
TOPIC = "酶"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组[｜|]\s*(.+)$")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.+)$")
ANSWER_RE = re.compile(r"(\d+)\.\s*([A-Z](?:、[A-Z])*)")


def parse_workbook():
    doc = Document(SOURCE)
    groups = []
    current = None
    mode = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        match = GROUP_RE.match(text)
        if match:
            current = {"source_index": int(match.group(1)), "title": match.group(2).strip(), "stems": [], "answers": {}}
            groups.append(current)
            mode = ""
            continue
        if current is None:
            continue
        if text == "题干":
            mode = "stems"
            continue
        if text.startswith("选项池"):
            mode = "options"
            continue
        if text == "答案":
            mode = "answers"
            continue
        if mode == "stems":
            question = QUESTION_RE.match(text)
            if question:
                current["stems"].append((int(question.group(1)), question.group(2).strip()))
        elif mode == "answers":
            for number, answer in ANSWER_RE.findall(text):
                current["answers"][int(number)] = answer.replace("、", "")

    option_banks = []
    for table in doc.tables:
        option_banks.append([(row.cells[0].text.strip(), row.cells[1].text.strip()) for row in table.rows[1:]])
    if len(groups) != len(option_banks):
        raise ValueError(f"Found {len(groups)} groups but {len(option_banks)} option banks")
    for group, options in zip(groups, option_banks):
        if set(group["answers"]) != {number for number, _ in group["stems"]}:
            raise ValueError(f"Group {group['source_index']}: question and answer keys do not match")
        group["options"] = options
    return groups


def evidence(page):
    return {
        "lectureId": "lecture-12",
        "lectureNumber": LECTURE_NUMBER,
        "lectureTitle": TITLE,
        "page": page,
        "image": f"biochemistry/lecture-pages/lecture-12-page-{page:02d}.webp",
        "title": f"第 12 讲《{TITLE}》· 第 {page} 页",
        "description": "已按该讲义页逐项核对答案；点击可查看讲义原页。",
        "method": "按 2027 考研生化第 12 讲及配套思维导图逐项复核。",
    }


def make_group(source_group, display_index, evidence_page):
    original = dict(source_group["options"])
    labels = list(original.values())
    shuffled = list(labels)
    random.Random(30600 + LECTURE_NUMBER * 100 + source_group["source_index"]).shuffle(shuffled)
    if shuffled == labels:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + position) for position, label in enumerate(shuffled)}
    stems = []
    for number, text in source_group["stems"]:
        answers = [output_keys[original[key]] for key in source_group["answers"][number]]
        stems.append({
            "number": number,
            "text": text.replace("（多选）", "").rstrip(),
            "answerRaw": "、".join(answers),
            "answer": answers,
            "answerMode": "多选" if len(answers) > 1 else "单选",
        })
    return {
        "id": f"bio-12-{display_index:02d}",
        "page": display_index,
        "title": source_group["title"],
        "kind": "B",
        "kindLabel": "B型题",
        "options": [{"key": chr(65 + position), "label": label} for position, label in enumerate(shuffled)],
        "stems": stems,
        "sourceText": source_group["title"],
        "reviewState": "已按 2027 考研讲义与思维导图核对",
        "reviewIssues": [],
        "reviewNotes": [],
        "topic": TOPIC,
        "lectureIds": ["lecture-12"],
        "optionShuffleVersion": 1,
        "lectureEvidence": evidence(evidence_page),
    }


def main():
    source_groups = parse_workbook()
    # Lecture pages: active centre (1), catalyst properties (2–3), regulation (3),
    # kinetics (4), inhibitors (5), and holoenzymes/cofactors (6).
    evidence_pages = {1: 1, 2: 3, 3: 3, 4: 4, 5: 5, 6: 5, 7: 5, 8: 6}
    groups = [make_group(group, index, evidence_pages[group["source_index"]]) for index, group in enumerate(source_groups, 1)]
    payload = {
        "meta": {
            "title": "生物化学第 12 讲题库",
            "sourceLabel": "生化第 12 讲学成选择题（酶）",
            "sourcePages": 1,
            "lectureCount": 1,
            "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups),
            "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture12.py",
            "siteIntegrated": True,
            "lectureLinked": True,
            "answerNote": "仅收录第 12 讲《酶》范围内题目；每组选择题选项均已重新打散，答案已按讲义与思维导图复核。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": "lecture-12", "number": LECTURE_NUMBER, "title": TITLE, "pageCount": 9}],
    }
    Path("src/data/biochemistry-lecture12-data.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
