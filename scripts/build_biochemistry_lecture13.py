#!/usr/bin/env python3
"""Convert the checked vitamin workbook into the lecture 13 site payload."""

from __future__ import annotations

import json
import random
import re
from pathlib import Path

from docx import Document


SOURCE = Path("/Users/ray/Downloads/生化_维生素_学成选择题_连续编号版.docx")
LECTURE_NUMBER = 13
TITLE = "生化 维生素"
TOPIC = "维生素"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组[｜|]\s*(.+)$")
ANSWER_GROUP_RE = re.compile(r"^第\s*(\d+)\s*组$")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.+)$")


def parse_options(table):
    return [
        (row.cells[0].text.strip(), row.cells[1].text.strip())
        for row in table.rows[1:]
        if row.cells[0].text.strip() and row.cells[1].text.strip()
    ]


def parse_workbook():
    doc = Document(SOURCE)
    groups = []
    current = None
    mode = ""
    answer_group = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "答案":
            mode = "answers"
            continue
        if mode == "answers":
            answer_match = ANSWER_GROUP_RE.match(text)
            if answer_match:
                answer_group = int(answer_match.group(1))
                continue
            question = QUESTION_RE.match(text)
            if question and answer_group is not None:
                groups[answer_group - 1]["answers"][int(question.group(1))] = re.findall(r"[A-Z]", question.group(2))
            continue

        group_match = GROUP_RE.match(text)
        if group_match:
            current = {
                "source_index": int(group_match.group(1)),
                "title": group_match.group(2).strip(),
                "stems": [],
                "answers": {},
            }
            groups.append(current)
            mode = ""
            continue
        if current is None:
            continue
        if text == "题目":
            mode = "stems"
            continue
        if mode == "stems":
            question = QUESTION_RE.match(text)
            if question:
                current["stems"].append((int(question.group(1)), question.group(2).strip()))

    # The source has one logical option pool per group, split across tables only
    # for readability. Keep the letters continuous after joining those pieces.
    table_counts = [2, 1, 2, 2, 3]
    tables = iter(doc.tables)
    for group, count in zip(groups, table_counts):
        options = []
        for _ in range(count):
            options.extend(parse_options(next(tables)))
        group["options"] = options
        expected = {number for number, _ in group["stems"]}
        if set(group["answers"]) != expected:
            raise ValueError(f"Group {group['source_index']}: question and answer keys do not match")
        if set(key for key, _ in options) != {chr(65 + i) for i in range(len(options))}:
            raise ValueError(f"Group {group['source_index']}: option keys are not continuous")
        for _, letters in group["answers"].items():
            if not set(letters).issubset(dict(options)):
                raise ValueError(f"Group {group['source_index']}: answer has an unknown option")
    try:
        next(tables)
    except StopIteration:
        return groups
    raise ValueError("Unexpected extra option table")


def evidence():
    return {
        "lectureId": "lecture-13",
        "lectureNumber": LECTURE_NUMBER,
        "lectureTitle": TITLE,
        "page": "思维导图",
        "image": "biochemistry/lecture-pages/lecture-13-mind-map.webp",
        "title": f"第 13 讲《{TITLE}》· 思维导图",
        "description": "已按维生素讲义、真题要点与辅因子表逐项核对；点击可放大查看思维导图。",
        "method": "按 2027 考研生化第 13 讲维生素思维导图及配套选择题逐项复核。",
    }


def make_group(source_group, display_index):
    original = dict(source_group["options"])
    labels = list(original.values())
    shuffled = list(labels)
    random.Random(30600 + LECTURE_NUMBER * 100 + source_group["source_index"]).shuffle(shuffled)
    if shuffled == labels:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + position) for position, label in enumerate(shuffled)}
    stems = []
    for number, text in source_group["stems"]:
        answers = [output_keys[original[key]] for key in source_group["answers"][number]]
        stems.append({
            "number": number,
            "text": text.replace("（多选）", "").rstrip(),
            "answerRaw": "、".join(answers),
            "answer": answers,
            "answerMode": "多选" if len(answers) > 1 else "单选",
        })
    return {
        "id": f"bio-13-{display_index:02d}",
        "page": display_index,
        "title": source_group["title"],
        "kind": "B",
        "kindLabel": "B型题",
        "options": [{"key": chr(65 + position), "label": label} for position, label in enumerate(shuffled)],
        "stems": stems,
        "sourceText": source_group["title"],
        "reviewState": "已按维生素讲义、真题要点与思维导图核对",
        "reviewIssues": [],
        "reviewNotes": [],
        "topic": TOPIC,
        "lectureIds": ["lecture-13"],
        "optionShuffleVersion": 1,
        "lectureEvidence": evidence(),
    }


def main():
    source_groups = parse_workbook()
    groups = [make_group(group, index) for index, group in enumerate(source_groups, 1)]
    payload = {
        "meta": {
            "title": "生物化学第 13 讲题库",
            "sourceLabel": "生化第 13 讲学成选择题（维生素）",
            "sourcePages": 1,
            "lectureCount": 1,
            "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups),
            "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture13.py",
            "siteIntegrated": True,
            "lectureLinked": True,
            "answerNote": "仅收录第 13 讲《维生素》范围内题目；每组选项均已打散，答案按讲义、真题与辅因子表复核。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": "lecture-13", "number": LECTURE_NUMBER, "title": TITLE, "pageCount": 2}],
    }
    Path("src/data/biochemistry-lecture13-data.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
