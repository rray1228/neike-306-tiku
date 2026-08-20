#!/usr/bin/env python3
"""Build the checked nucleic-acid question bank payload."""

from __future__ import annotations

import json
import random
import re
from pathlib import Path

from docx import Document

from biochemistry_fill_questions import make_fill_groups, parse_fill_questions


SOURCE = Path("/Users/ray/Downloads/生化_15_核酸_学成题_讲义原话修订版.docx")
OUTPUT = Path("src/data/biochemistry-lecture15-data.json")
LECTURE_NUMBER = 15
TITLE = "生化 核酸"
TOPIC = "核酸"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组[｜|]\s*(.+)$")
ANSWER_GROUP_RE = re.compile(r"^第\s*(\d+)\s*组$")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.+)$")
ANSWER_RE = re.compile(r"(\d+)\.\s*([A-Z](?:、[A-Z])*)")


def parse_workbook():
    document = Document(SOURCE)
    groups = []
    current = None
    mode = ""
    answer_group = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "答案":
            mode = "answers"
            continue
        if mode == "answers":
            if text == "填空题答案":
                break
            answer_group_match = ANSWER_GROUP_RE.match(text)
            if answer_group_match:
                answer_group = int(answer_group_match.group(1))
                continue
            if answer_group is not None:
                for number, answer in ANSWER_RE.findall(text):
                    groups[answer_group - 1]["answers"][int(number)] = list(answer.replace("、", ""))
            continue

        group_match = GROUP_RE.match(text)
        if group_match:
            current = {
                "source_index": int(group_match.group(1)),
                "title": group_match.group(2).strip(),
                "stems": [],
                "answers": {},
            }
            groups.append(current)
            mode = ""
            continue
        if current is not None and text == "题目":
            mode = "stems"
            continue
        if mode == "stems":
            question_match = QUESTION_RE.match(text)
            if question_match:
                current["stems"].append((int(question_match.group(1)), question_match.group(2).strip()))

    option_banks = [
        [(row.cells[0].text.strip(), row.cells[1].text.strip()) for row in table.rows[1:]]
        for table in document.tables
    ]
    if len(groups) != len(option_banks):
        raise ValueError(f"Found {len(groups)} groups but {len(option_banks)} option banks")
    for group, options in zip(groups, option_banks):
        expected_numbers = {number for number, _ in group["stems"]}
        if set(group["answers"]) != expected_numbers:
            raise ValueError(f"Group {group['source_index']}: question and answer keys do not match")
        option_map = dict(options)
        if not all(set(answer).issubset(option_map) for answer in group["answers"].values()):
            raise ValueError(f"Group {group['source_index']}: answer has an unknown option")
        group["options"] = options
    return groups


def evidence(page):
    return {
        "lectureId": "lecture-15",
        "lectureNumber": LECTURE_NUMBER,
        "lectureTitle": TITLE,
        "page": page,
        "image": f"biochemistry/lecture-pages/lecture-15-page-{page:02d}.webp",
        "title": f"第 15 讲《{TITLE}》· 第 {page} 页",
        "description": "已按核酸讲义、思维导图与生化合集逐项核对；点击可查看讲义原页。",
        "method": "按 2027 考研生化第 15 讲《核酸》及配套思维导图逐项复核。",
    }


def make_group(source_group, display_index, evidence_page):
    original = dict(source_group["options"])
    labels = list(original.values())
    shuffled = labels.copy()
    random.Random(30600 + LECTURE_NUMBER * 100 + source_group["source_index"]).shuffle(shuffled)
    if shuffled == labels and len(shuffled) > 1:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + index) for index, label in enumerate(shuffled)}
    stems = []
    for number, text in source_group["stems"]:
        answers = [output_keys[original[key]] for key in source_group["answers"][number]]
        stems.append({
            "number": number,
            "text": text.replace("（多选）", "").rstrip(),
            "answerRaw": "、".join(answers),
            "answer": answers,
            "answerMode": "多选" if len(answers) > 1 else "单选",
        })
    return {
        "id": f"bio-15-{display_index:02d}",
        "page": display_index,
        "title": source_group["title"],
        "kind": "B",
        "kindLabel": "B型题",
        "options": [{"key": chr(65 + index), "label": label} for index, label in enumerate(shuffled)],
        "stems": stems,
        "sourceText": source_group["title"],
        "reviewState": "已按核酸讲义、思维导图与生化合集核对",
        "reviewIssues": [],
        "reviewNotes": [],
        "topic": TOPIC,
        "lectureIds": ["lecture-15"],
        "optionShuffleVersion": 1,
        "lectureEvidence": evidence(evidence_page),
    }


def main():
    source_groups = parse_workbook()
    evidence_pages = {**{index: 1 for index in range(1, 5)}, **{index: 2 for index in range(5, 8)}}
    groups = [make_group(group, index, evidence_pages[group["source_index"]]) for index, group in enumerate(source_groups, 1)]
    fill_questions = parse_fill_questions(SOURCE)
    groups.extend(make_fill_groups(
        fill_questions,
        lecture_number=LECTURE_NUMBER,
        topic=TOPIC,
        lecture_id="lecture-15",
        start_page=len(groups) + 1,
        ranges=[(1, 7, 1), (8, 13, 1), (14, 19, 2), (20, 25, 2)],
        evidence_for_page=evidence,
        review_state="已按核酸 DOCX 原文逐空核对",
    ))
    payload = {
        "meta": {
            "title": "生物化学第 15 讲题库",
            "sourceLabel": "生化第 15 讲学成题（核酸）",
            "sourcePages": 2,
            "lectureCount": 1,
            "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups),
            "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture15.py",
            "siteIntegrated": True,
            "lectureLinked": True,
            "answerNote": "完整收录第 15 讲《核酸》选择题与填空题；选择题选项已打散，填空题按 DOCX 原编号与原答案导入。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": "lecture-15", "number": LECTURE_NUMBER, "title": TITLE, "pageCount": 2}],
    }
    OUTPUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
