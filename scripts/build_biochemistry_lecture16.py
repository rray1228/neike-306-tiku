#!/usr/bin/env python3
"""Build the checked DNA-synthesis question bank and linked lecture images."""

from __future__ import annotations

import json
import random
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from PIL import Image

from biochemistry_fill_questions import make_fill_groups, parse_fill_questions


SOURCE = Path("/Users/ray/Downloads/生化_16_DNA的合成_学成题_讲义原话版.docx")
LECTURE_PDF = Path("/Users/ray/Desktop/306/生物化学/讲义/16 27考研：生化 DNA的合成 核心-真题-串联-导图 天天师兄.pdf")
OUTPUT = Path("src/data/biochemistry-lecture16-data.json")
IMAGE_DIR = Path("public/biochemistry/lecture-pages")
LECTURE_NUMBER = 16
TITLE = "生化 DNA 的合成"
TOPIC = "核酸"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组[｜|]\s*(.+)$")
ANSWER_GROUP_RE = re.compile(r"^第\s*(\d+)\s*组(?:\s+|$)")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.+)$")
ANSWER_RE = re.compile(r"(\d+)\.\s*([A-Z](?:、[A-Z])*)")

# This second, fixed check catches any accidental mismatch while importing the DOCX.
EXPECTED_ANSWER_KEYS = {
    1: {1: "BDF", 2: "E", 3: "C"},
    2: {1: "BCEGHI"},
    3: {1: "C", 2: "A", 3: "D", 4: "B"},
    4: {1: "D", 2: "B", 3: "C"},
    5: {1: "E", 2: "F", 3: "C", 4: "G", 5: "D", 6: "B", 7: "A"},
    6: {1: "C", 2: "B", 3: "A"},
    7: {1: "D", 2: "B", 3: "AC"},
    8: {1: "B", 2: "D", 3: "A", 4: "E", 5: "C"},
    9: {1: "BDF"},
    10: {1: "BCD"},
    11: {1: "C"},
    12: {1: "ACDFGHI"},
    13: {1: "ACD"},
}


def parse_workbook():
    document = Document(SOURCE)
    groups = []
    current = None
    mode = ""
    answer_group = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "答案":
            mode = "answers"
            continue
        if mode == "answers":
            if text == "填空题答案":
                break
            answer_group_match = ANSWER_GROUP_RE.match(text)
            if answer_group_match:
                answer_group = int(answer_group_match.group(1))
                text = text[answer_group_match.end():].strip()
            if answer_group is not None:
                for number, answer in ANSWER_RE.findall(text):
                    groups[answer_group - 1]["answers"][int(number)] = list(answer.replace("、", ""))
            continue

        group_match = GROUP_RE.match(text)
        if group_match:
            current = {
                "source_index": int(group_match.group(1)),
                "title": group_match.group(2).strip(),
                "stems": [],
                "answers": {},
            }
            groups.append(current)
            mode = ""
            continue
        if current is not None and text == "题目":
            mode = "stems"
            continue
        if mode == "stems":
            question_match = QUESTION_RE.match(text)
            if question_match:
                current["stems"].append((int(question_match.group(1)), question_match.group(2).strip()))

    option_banks = [
        [(row.cells[0].text.strip(), row.cells[1].text.strip()) for row in table.rows[1:]]
        for table in document.tables
    ]
    if len(groups) != len(option_banks) or len(groups) != 13:
        raise ValueError(f"Found {len(groups)} groups and {len(option_banks)} option banks")
    for group, options in zip(groups, option_banks):
        expected_numbers = {number for number, _ in group["stems"]}
        if set(group["answers"]) != expected_numbers:
            raise ValueError(f"Group {group['source_index']}: question and answer keys do not match")
        answer_key_check = {number: "".join(answer) for number, answer in group["answers"].items()}
        if answer_key_check != EXPECTED_ANSWER_KEYS[group["source_index"]]:
            raise ValueError(f"Group {group['source_index']}: answer key differs from checked import")
        option_map = dict(options)
        if not all(set(answer).issubset(option_map) for answer in group["answers"].values()):
            raise ValueError(f"Group {group['source_index']}: answer has an unknown option")
        group["options"] = options
    return groups


def find_pdftoppm():
    return (
        shutil.which("pdftoppm")
        or "/Users/ray/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override/pdftoppm"
    )


def render_pdf_page(page_number: int, output: Path):
    if not LECTURE_PDF.is_file():
        raise FileNotFoundError(f"Missing checked lecture PDF: {LECTURE_PDF}")
    output.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            find_pdftoppm(), "-f", str(page_number), "-l", str(page_number),
            "-singlefile", "-png", "-r", "150", str(LECTURE_PDF), str(output.with_suffix("")),
        ],
        check=True,
    )


def save_webp(image: Image.Image, output: Path):
    output.parent.mkdir(parents=True, exist_ok=True)
    image.convert("RGB").save(output, "WEBP", quality=88, method=6, exact=True)


def build_lecture_images():
    """Keep only mapped lecture pages; group 1 requires both pages 1 and 2."""
    with tempfile.TemporaryDirectory(prefix="bio-lecture16-") as temporary_directory:
        temporary = Path(temporary_directory)
        pages = {}
        for page_number in (1, 2, 3, 4, 5, 8, 9, 10, 11):
            rendered = temporary / f"page-{page_number:02d}.png"
            render_pdf_page(page_number, rendered)
            pages[page_number] = Image.open(rendered).convert("RGB")

        width = max(pages[1].width, pages[2].width)
        gap = 24
        combined = Image.new("RGB", (width, pages[1].height + gap + pages[2].height), "white")
        combined.paste(pages[1], ((width - pages[1].width) // 2, 0))
        combined.paste(pages[2], ((width - pages[2].width) // 2, pages[1].height + gap))
        save_webp(combined, IMAGE_DIR / "lecture-16-page-01-02.webp")

        for page_number in (3, 4, 5, 8, 9, 10, 11):
            save_webp(pages[page_number], IMAGE_DIR / f"lecture-16-page-{page_number:02d}.webp")


def evidence(source_group):
    evidence_pages = {
        1: ("01-02", "第 1–2 页", "DNA 复制方式：半保留、双向、半不连续与多复制子。"),
        2: ("03", "第 3 页", "DNA 复制高保真性及原核 DNA 聚合酶。"),
        3: ("03", "第 3 页", "原核 DNA 聚合酶 I、II、III 的功能。"),
        4: ("04", "第 4 页", "DNA 聚合酶 III 的亚基、复制相关酶与蛋白。"),
        5: ("04", "第 4 页", "DnaA、DnaB、DnaC、DnaG、SSB 与拓扑异构酶。"),
        6: ("04", "第 4 页", "前导链、后随链与冈崎片段。"),
        7: ("05", "第 5 页", "RNA 引物切除、填补与 DNA 连接。"),
        8: ("08", "第 8 页", "真核 DNA 聚合酶及线粒体 DNA 的 D 环复制。"),
        9: ("09", "第 9 页", "端粒酶的组成。"),
        10: ("10", "第 10 页", "逆转录酶的三种酶活性。"),
        11: ("11", "第 11 页", "端粒酶具有逆转录活性。"),
        12: ("10", "第 10 页", "DNA、RNA 合成及拓扑反应中的磷酸二酯键生成。"),
        13: ("11", "第 11 页", "DNA 连接酶在复制、修复及基因工程中的作用。"),
    }
    image_page, page_title, description = evidence_pages[source_group]
    return {
        "lectureId": "lecture-16",
        "lectureNumber": LECTURE_NUMBER,
        "lectureTitle": TITLE,
        "page": source_group,
        "image": f"biochemistry/lecture-pages/lecture-16-page-{image_page}.webp",
        "title": f"第 16 讲《{TITLE}》· {page_title}",
        "description": f"{description} 点击可查看讲义原页。",
        "method": "按第 16 讲《DNA 的合成》原始讲义逐项复核。",
    }


def make_group(source_group, display_index):
    original = dict(source_group["options"])
    labels = list(original.values())
    shuffled = labels.copy()
    random.Random(30600 + LECTURE_NUMBER * 100 + source_group["source_index"]).shuffle(shuffled)
    if shuffled == labels and len(shuffled) > 1:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + index) for index, label in enumerate(shuffled)}
    stems = []
    for number, text in source_group["stems"]:
        answers = [output_keys[original[key]] for key in source_group["answers"][number]]
        stems.append({
            "number": number,
            "text": text.replace("（多选）", "").rstrip(),
            "answerRaw": "、".join(answers),
            "answer": answers,
            "answerMode": "多选" if len(answers) > 1 else "单选",
        })
    return {
        "id": f"bio-16-{display_index:02d}",
        "page": display_index,
        "title": source_group["title"],
        "kind": "B",
        "kindLabel": "B型题",
        "options": [{"key": chr(65 + index), "label": label} for index, label in enumerate(shuffled)],
        "stems": stems,
        "sourceText": source_group["title"],
        "reviewState": "已按 DNA 的合成原始讲义核对",
        "reviewIssues": [],
        "reviewNotes": [],
        "topic": TOPIC,
        "lectureIds": ["lecture-16"],
        "optionShuffleVersion": 1,
        "lectureEvidence": evidence(source_group["source_index"]),
    }


def main():
    source_groups = parse_workbook()
    build_lecture_images()
    groups = [make_group(group, index) for index, group in enumerate(source_groups, 1)]
    fill_questions = parse_fill_questions(SOURCE)
    groups.extend(make_fill_groups(
        fill_questions,
        lecture_number=LECTURE_NUMBER,
        topic=TOPIC,
        lecture_id="lecture-16",
        start_page=len(groups) + 1,
        ranges=[(1, 5, 2), (6, 11, 5), (12, 12, 7), (13, 16, 8), (17, 18, 9), (19, 23, 10)],
        evidence_for_page=lambda source_group: evidence(source_group),
        review_state="已按 DNA 的合成 DOCX 原文逐空核对",
    ))
    payload = {
        "meta": {
            "title": "生物化学第 16 讲题库",
            "sourceLabel": "生化第 16 讲学成题（DNA 的合成）",
            "sourcePages": 13,
            "lectureCount": 1,
            "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups),
            "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture16.py",
            "siteIntegrated": True,
            "lectureLinked": True,
            "answerNote": "完整收录第 16 讲《DNA 的合成》选择题与填空题；选择题选项已打散，填空题按 DOCX 原编号与原答案导入。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": "lecture-16", "number": LECTURE_NUMBER, "title": TITLE, "pageCount": 13}],
    }
    OUTPUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
