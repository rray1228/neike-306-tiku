#!/usr/bin/env python3
"""Convert the supplied, checked lecture 07/08 workbooks into site payloads."""

from __future__ import annotations

import json
import random
import re
from pathlib import Path

from docx import Document


SOURCE_DIR = Path("/Users/ray/Downloads")
TOPIC = "氨基酸与蛋白质"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组\s+(.*)$")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.*)$")
AMINO_ACID_ABBREVIATIONS = {
    "Valine": "Val", "Glutamine": "Gln", "Phenylalanine": "Phe", "Threonine": "Thr",
    "Proline": "Pro", "Cysteine": "Cys", "Arginine": "Arg", "Leucine": "Leu",
    "Tryptophan": "Trp", "Histidine": "His", "Aspartic acid": "Asp", "Glutamic acid": "Glu",
    "Asparagine": "Asn", "Alanine": "Ala", "Serine": "Ser", "Glycine": "Gly",
    "Tyrosine": "Tyr", "Methionine": "Met", "Lysine": "Lys", "Isoleucine": "Ile",
}


def parse_workbook(path: Path):
    doc = Document(path)
    source_groups = []
    answers = {}
    current = None
    in_answers = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "参考答案":
            in_answers = True
            current = None
            continue
        group_match = GROUP_RE.match(text)
        if group_match:
            current = int(group_match.group(1))
            title = group_match.group(2).strip()
            if in_answers:
                answers[current] = []
            else:
                source_groups.append({"source_index": current, "title": title, "stems": []})
            continue
        question_match = QUESTION_RE.match(text)
        if not question_match or current is None:
            continue
        number = int(question_match.group(1))
        content = question_match.group(2).strip()
        if in_answers:
            answers[current].append((number, content.replace("、", "")))
        else:
            source_groups[-1]["stems"].append((number, content))
    option_banks = []
    for table in doc.tables:
        options = []
        for row in table.rows[1:]:
            key, label = (cell.text.strip() for cell in row.cells[:2])
            options.append((key, label))
        option_banks.append(options)
    if len(option_banks) != len(source_groups):
        raise ValueError(f"{path.name}: found {len(source_groups)} groups but {len(option_banks)} option banks")
    for group, options in zip(source_groups, option_banks):
        source_answer_rows = answers.get(group["source_index"], [])
        if len(source_answer_rows) != len(group["stems"]):
            raise ValueError(f"{path.name} group {group['source_index']}: question/answer count mismatch")
        group["options"] = options
        group["answers"] = {number: list(answer) for number, answer in source_answer_rows}
    return source_groups


def evidence(lecture_number, lecture_title, page):
    return {
        "lectureId": f"lecture-{lecture_number:02d}", "lectureNumber": lecture_number,
        "lectureTitle": lecture_title, "page": page,
        "image": f"biochemistry/lecture-pages/lecture-{lecture_number:02d}-page-{page:02d}.webp",
        "title": f"第 {lecture_number:02d} 讲《{lecture_title}》· 第 {page} 页",
        "description": "已按该讲义页逐项核对答案；点击可查看讲义原页。",
        "method": f"按知识点人工映射至 2027 考研生化第 {lecture_number:02d} 讲，并逐项复核。",
    }


def make_group(lecture_number, lecture_title, display_index, source_group, evidence_page):
    original = list(source_group["options"])
    use_abbreviations = lecture_number == 7 and source_group["source_index"] == 11
    source_labels = {
        key: AMINO_ACID_ABBREVIATIONS.get(label, label) if use_abbreviations else label
        for key, label in original
    }
    option_labels = list(source_labels.values())
    shuffled = list(option_labels)
    random.Random(30600 + lecture_number * 100 + source_group["source_index"]).shuffle(shuffled)
    if shuffled == option_labels:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + position) for position, label in enumerate(shuffled)}
    stems = []
    for number, text in source_group["stems"]:
        raw_answers = source_group["answers"][number]
        try:
            answer = [output_keys[source_labels[key]] for key in raw_answers]
        except KeyError as error:
            raise ValueError(f"Lecture {lecture_number} group {source_group['source_index']} stem {number}: unknown answer key {error}") from error
        stems.append({"number": number, "text": text.replace("（多选）", "").rstrip(), "answerRaw": "、".join(answer), "answer": answer, "answerMode": "多选" if len(answer) > 1 else "单选"})
    return {
        "id": f"bio-{lecture_number:02d}-{display_index:02d}", "page": display_index,
        "title": source_group["title"], "kind": "B", "kindLabel": "B型题",
        "options": [{"key": chr(65 + position), "label": label} for position, label in enumerate(shuffled)],
        "stems": stems, "sourceText": source_group["title"], "reviewState": "已按 2027 考研讲义核对",
        "reviewIssues": [], "reviewNotes": [], "topic": TOPIC,
        "lectureIds": [f"lecture-{lecture_number:02d}"], "optionShuffleVersion": 1,
        "lectureEvidence": evidence(lecture_number, lecture_title, evidence_page),
    }


def make_custom_group(lecture_number, lecture_title, display_index, title, options, stems, evidence_page):
    shuffled = list(options)
    random.Random(30600 + lecture_number * 100 + display_index).shuffle(shuffled)
    if shuffled == options:
        shuffled = shuffled[1:] + shuffled[:1]
    output_keys = {label: chr(65 + position) for position, label in enumerate(shuffled)}
    return {
        "id": f"bio-{lecture_number:02d}-{display_index:02d}", "page": display_index,
        "title": title, "kind": "B", "kindLabel": "B型题",
        "options": [{"key": chr(65 + position), "label": label} for position, label in enumerate(shuffled)],
        "stems": [{
            "number": number, "text": text, "answerRaw": "、".join(output_keys[item] for item in answer),
            "answer": [output_keys[item] for item in answer], "answerMode": "多选",
        } for number, (text, answer) in enumerate(stems, 1)],
        "sourceText": title, "reviewState": "已按 2027 考研讲义核对", "reviewIssues": [], "reviewNotes": [],
        "topic": TOPIC, "lectureIds": [f"lecture-{lecture_number:02d}"], "optionShuffleVersion": 1,
        "lectureEvidence": evidence(lecture_number, lecture_title, evidence_page),
    }


def payload(lecture_number, lecture_title, source_name, groups):
    return {
        "meta": {
            "title": f"生物化学第 {lecture_number:02d} 讲题库", "sourceLabel": source_name,
            "sourcePages": 1, "lectureCount": 1, "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups), "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture7_8.py", "siteIntegrated": True,
            "lectureLinked": True, "answerNote": f"仅收录第 {lecture_number:02d} 讲《{lecture_title}》范围内题目；选项已逐组打散，答案已按讲义复核。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": f"lecture-{lecture_number:02d}", "number": lecture_number, "title": lecture_title, "pageCount": 0}],
    }


def main():
    title7 = "生化 氨基酸代谢与氨基酸"
    groups7 = parse_workbook(SOURCE_DIR / "生化_07_氨基酸代谢与氨基酸_学成选择题_中英文名补充版.docx")
    # User-specified order: the original group 11 is the first study group.
    groups7 = [next(group for group in groups7 if group["source_index"] == 11)] + [group for group in groups7 if group["source_index"] != 11]
    groups7[0]["title"] = "氨基酸中文名与英文缩写"
    evidence7 = {11: 7, 1: 7, 2: 1, 3: 2, 4: 2, 5: 2, 6: 3, 7: 3, 8: 4, 9: 1, 10: 1}
    output7 = [make_group(7, title7, index, group, evidence7[group["source_index"]]) for index, group in enumerate(groups7, 1)]
    data7 = payload(7, title7, "生化第 07 讲学成选择题（氨基酸代谢与氨基酸）", output7)
    data7["lectures"][0]["pageCount"] = 12
    Path("src/data/biochemistry-lecture7-data.json").write_text(json.dumps(data7, ensure_ascii=False, indent=2), encoding="utf-8")

    title8 = "生化 蛋白质"
    groups8 = parse_workbook(SOURCE_DIR / "生化_08_蛋白质_学成选择题_选项打散版.docx")
    # 保留原第 05、06 组题目。
    # 原第 06 组的若干选项把完整结论直接写出来（如“白蛋白最快”），
    # 会在题干尚未作答时泄露答案。将复合结论拆为中性知识点，并把
    # “最快 / 最慢”分别改为独立题干，答案仍逐项对应讲义。
    chromatography_group = next(group for group in groups8 if group["source_index"] == 6)
    chromatography_group["title"] = "层析、电泳与等电点"
    chromatography_group["options"] = [
        ("A", "白蛋白"), ("B", "分子量小"), ("C", "溶液离子强度小"),
        ("D", "电荷量大"), ("E", "电中性"), ("F", "γ球蛋白"),
        ("G", "分子量大"), ("H", "电荷量小"), ("I", "带负电"),
        ("J", "分子量"), ("K", "带正电"), ("L", "球状"),
    ]
    chromatography_group["stems"] = [
        (1, "凝胶过滤 / 分子筛层析中先洗脱的分子特点"),
        (2, "离子交换层析中先被分离出的分子特点"),
        (3, "电泳中跑得较快的因素"),
        (4, "血清蛋白电泳中跑得最快的蛋白"),
        (5, "血清蛋白电泳中跑得最慢的蛋白"),
        (6, "SDS-PAGE 的主要分离依据"),
        (7, "pH＜pI 时蛋白质所带电荷"),
        (8, "pH＝pI 时蛋白质所带电荷"),
        (9, "pH＞pI 时蛋白质所带电荷"),
    ]
    chromatography_group["answers"] = {
        1: list("G"), 2: list("H"), 3: list("BLCD"), 4: list("A"), 5: list("F"),
        6: list("J"), 7: list("K"), 8: list("E"), 9: list("I"),
    }
    output8 = [make_group(8, title8, index, group, 1) for index, group in enumerate(groups8, 1)]
    data8 = payload(8, title8, "生化第 08 讲学成选择题（蛋白质）", output8)
    data8["lectures"][0]["pageCount"] = 5
    Path("src/data/biochemistry-lecture8-data.json").write_text(json.dumps(data8, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
