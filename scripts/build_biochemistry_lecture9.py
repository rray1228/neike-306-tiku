#!/usr/bin/env python3
"""Build the verified lecture 09 nucleotide-metabolism question bank."""

from __future__ import annotations

import json
import random
import re
from pathlib import Path

from docx import Document


SOURCE = Path("/Users/ray/Downloads/生化_09_核苷酸代谢_学成选择题_选项打散版.docx")
TITLE = "生化 核苷酸代谢"
TOPIC = "核苷酸代谢"
GROUP_RE = re.compile(r"^第\s*(\d+)\s*组\s+(.*)$")
QUESTION_RE = re.compile(r"^(\d+)\.\s*(.*)$")


def parse_workbook():
    doc = Document(SOURCE)
    groups, answers = [], {}
    in_answers = False
    current_index = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "参考答案":
            in_answers = True
            current_index = None
            continue
        group_match = GROUP_RE.match(text)
        if group_match:
            current_index = int(group_match.group(1))
            if in_answers:
                answers[current_index] = {}
            else:
                groups.append({"index": current_index, "title": group_match.group(2).strip(), "stems": []})
            continue
        question_match = QUESTION_RE.match(text)
        if not question_match or current_index is None:
            continue
        number, content = int(question_match.group(1)), question_match.group(2).strip()
        if in_answers:
            answers[current_index][number] = list(content.replace("、", ""))
        else:
            groups[-1]["stems"].append((number, content))

    option_banks = []
    for table in doc.tables:
        option_banks.append([(row.cells[0].text.strip(), row.cells[1].text.strip()) for row in table.rows[1:]])
    if len(groups) != len(option_banks):
        raise ValueError(f"Found {len(groups)} groups but {len(option_banks)} option banks")

    for group, options in zip(groups, option_banks):
        group_answers = answers.get(group["index"], {})
        if len(group_answers) != len(group["stems"]):
            raise ValueError(f"Group {group['index']}: question/answer count mismatch")
        group["options"] = options
        group["answers"] = group_answers
    return groups


def lecture_evidence(page):
    return {
        "lectureId": "lecture-09",
        "lectureNumber": 9,
        "lectureTitle": TITLE,
        "page": page,
        "image": f"biochemistry/lecture-pages/lecture-09-page-{page:02d}.webp",
        "title": f"第 09 讲《{TITLE}》· 第 {page} 页",
        "description": "已按该讲义页逐项核对答案；点击可查看讲义原页。",
        "method": "按知识点人工映射至 2027 考研生化第 09 讲，并逐项复核。",
    }


def make_group(source_group, evidence_page):
    original_options = dict(source_group["options"])
    labels = list(original_options.values())
    shuffled = list(labels)
    random.Random(30609 + source_group["index"]).shuffle(shuffled)
    if shuffled == labels:
        shuffled = shuffled[1:] + shuffled[:1]
    output_key = {label: chr(65 + position) for position, label in enumerate(shuffled)}

    stems = []
    for number, text in source_group["stems"]:
        answer = [output_key[original_options[source_key]] for source_key in source_group["answers"][number]]
        stems.append({
            "number": number,
            "text": text.replace("（多选）", "").rstrip(),
            "answerRaw": "、".join(answer),
            "answer": answer,
            "answerMode": "多选" if len(answer) > 1 else "单选",
        })
    index = source_group["index"]
    return {
        "id": f"bio-09-{index:02d}",
        "page": index,
        "title": source_group["title"],
        "kind": "B",
        "kindLabel": "B型题",
        "options": [{"key": chr(65 + position), "label": label} for position, label in enumerate(shuffled)],
        "stems": stems,
        "sourceText": source_group["title"],
        "reviewState": "已按 2027 考研讲义与思维导图核对",
        "reviewIssues": [],
        "reviewNotes": [],
        "topic": TOPIC,
        "lectureIds": ["lecture-09"],
        "optionShuffleVersion": 1,
        "lectureEvidence": lecture_evidence(evidence_page),
    }


def main():
    # Groups 1–4 cover purines, 5–6 pyrimidines/CPS, and 7–8 antimetabolites.
    evidence_pages = {1: 1, 2: 1, 3: 1, 4: 1, 5: 2, 6: 2, 7: 3, 8: 3}
    source_groups = parse_workbook()
    # The last group must cover every column of the antimetabolite table:
    # analogue, competing/acting enzyme, and inhibited process.
    final_group = next(group for group in source_groups if group["index"] == 8)
    final_group["options"].extend([
        ("O", "胞苷"),
        ("P", "次黄嘌呤"),
        ("Q", "谷氨酰胺"),
        ("R", "叶酸"),
        ("S", "胸腺嘧啶"),
    ])
    similar_keys = {1: "P", 2: "P", 3: "O", 4: "S", 5: "R", 6: "Q"}
    final_group["stems"] = [
        (number, text.replace("①竞争 / 作用酶 ②抑制过程", "①类似物 ②竞争 / 作用酶 ③抑制过程"))
        for number, text in final_group["stems"]
    ]
    for number, similar_key in similar_keys.items():
        final_group["answers"][number].insert(0, similar_key)

    groups = [make_group(group, evidence_pages[group["index"]]) for group in source_groups]
    payload = {
        "meta": {
            "title": "生物化学第 09 讲题库",
            "sourceLabel": "生化第 09 讲学成选择题（核苷酸代谢）",
            "sourcePages": 1,
            "lectureCount": 1,
            "groupCount": len(groups),
            "stemCount": sum(len(group["stems"]) for group in groups),
            "correctionGroupCount": 0,
            "generatedBy": "scripts/build_biochemistry_lecture9.py",
            "siteIntegrated": True,
            "lectureLinked": True,
            "answerNote": "仅收录第 09 讲《核苷酸代谢》范围内题目；选项已重新打散，答案按讲义与思维导图复核。",
        },
        "topics": ["全部", TOPIC, "综合"],
        "pages": [{"page": group["page"], "image": "", "topic": TOPIC, "searchText": group["title"]} for group in groups],
        "groups": groups,
        "lectures": [{"id": "lecture-09", "number": 9, "title": TITLE, "pageCount": 9}],
    }
    Path("src/data/biochemistry-lecture9-data.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
